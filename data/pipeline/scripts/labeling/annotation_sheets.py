"""Module for generating and ingesting annotation sheets."""

import collections
import docx
import json
import hashlib

from pathlib import Path
from abc import ABC, abstractmethod
from typing import List, TypedDict, Optional, Union, Dict, Any, Tuple
from tf.fabric import Fabric
from docx.shared import Pt, RGBColor
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from labeling.specifiers import LingLabel, NodeIdentifier


# define some constants
GRAY = RGBColor(130, 130, 130)
RED = RGBColor(255, 0, 0)
SHEBANQ_LINK = (
    'https://shebanq.ancient-data.org/hebrew/text'
    '?book={book}&chapter={chapter}&verse={verse}&version=2021'
)


class AnnotationSheetSpecs(TypedDict):
    """TypedDict to hold specs for an annotation sheet."""
    sheet: str
    project: str


def add_hyperlink(paragraph, url, text, color=None, underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    Source: https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :param color: a color for the URL
    :param underline: whether to underline the URL
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


class BaseAnnotationSheet(ABC):
    """Object for generating an annotation sheet as a MS document."""

    NAME = "base"
    HASH_LEN = 8

    def __init__(
            self,
            annotations: List[LingLabel],
            tf_fabric: Fabric,
            project: 'BaseLabelingProject',
            document: Optional[Document] = None,
    ):
        """Initialize an AnnotationSheet object."""
        self.annotations = annotations
        self.annotation_metadata, self.label_to_id = (
            self._get_annotation_metadata(annotations)
        )
        self.project = project
        self.tf_fabric = tf_fabric
        self.tf_api = tf_fabric.api
        self.styles = {}
        if not document:
            self.document = Document()
            self._inject_specs_into_docx_metadata(self.document)
            self._add_styles()
            self._build_document(self.document)
        else:
            self.document = document

    @property
    def specs(self) -> AnnotationSheetSpecs:
        """Retrieve specs for this sheet template."""
        return {
            "sheet": self.NAME,
            "project": self.project.name,
        }

    def _inject_specs_into_docx_metadata(self, doc: Document):
        """Add annotation specs to the document's "comments" metadata section."""
        doc.core_properties.comments = json.dumps(self.specs)

    def to_docx(self, filepath: Path) -> None:
        """Save a new docx with document."""
        self.document.save(filepath)

    @abstractmethod
    def _add_styles(self) -> None:
        """Set all styles for the document."""

    @abstractmethod
    def _build_document(self, document: Document) -> None:
        """Build up document."""

    @staticmethod
    @abstractmethod
    def _label_from_row(row, metadata: Dict[str, Any]) -> LingLabel:
        """Extract cell values from a table row into a LingLabel object."""

    def _get_new_hash_id(
            self,
            label: LingLabel,
            annotation_metadata: Dict[str, Dict[str, Any]],
    ):
        """Get shortened annotation hash ID."""
        # turn label into string and hash it
        label_str = (
            label.label + str(label.nid) + label.target
        )
        label_hash = hashlib.sha1(label_str.encode()).hexdigest()

        # get unique truncated hash to prevent collisions
        trunc_len = self.HASH_LEN
        truncated_hash = label_hash[:trunc_len]
        while truncated_hash in annotation_metadata:
            trunc_len += 1
            truncated_hash = label_hash[:trunc_len]
        return truncated_hash

    def _get_annotation_metadata(
            self,
            labels: List[LingLabel]
    ) -> Tuple[Dict[str, Dict[str, Any]], Dict[LingLabel, str]]:
        """Store metadata for an annotation."""
        annotation_metadata = {}
        label_to_id = {}
        for label in labels:
            hash_id = self._get_new_hash_id(label, annotation_metadata)
            metadata = {
                'label': label.label,
                'nid': tuple(label.nid),
                'target': label.target
            }
            annotation_metadata[hash_id] = metadata
            label_to_id[label] = hash_id
        return annotation_metadata, label_to_id

    @classmethod
    def from_doc(
            cls,
            document: Union[Path, Document],
            tf_fabric: Fabric,
            project: 'BaseLabelingProject',
            metadata: Dict[str, Any],
    ) -> 'BaseAnnotationSheet':
        """Read in docx annotation sheet."""
        if isinstance(document, Path):
            document = Document(document)

        # populate annotations from docx
        annotations = []
        for table in document.tables:
            for row in table.rows:
                annotations.append(
                    cls._label_from_row(row, metadata)
                )

        # return new class instance
        return cls(
            annotations=annotations,
            tf_fabric=tf_fabric,
            project=project,
            document=document,
        )


class BasicAnnotationSheet(BaseAnnotationSheet):
    """Prepare an annotation sheet."""

    NAME = "basic"

    def _add_styles(self) -> None:
        """Set styles for the basic annotation sheet."""
        styles = self.styles
        doc_styles = self.document.styles

        # set reference header style
        styles['ref'] = doc_styles.add_style('Reference', WD_STYLE_TYPE.PARAGRAPH)
        styles['ref'].font.name = 'Times New Roman'
        styles['ref'].font.bold = True
        styles['ref'].font.size = Pt(12)
        styles['ref'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        styles['ref'].paragraph_format.keep_with_next = True

        # set hebrew text style
        styles['hebrew'] = doc_styles.add_style('Hebrew', WD_STYLE_TYPE.PARAGRAPH)
        styles['hebrew'].font.size = Pt(15)
        styles['hebrew'].font.name = 'SBL BibLit'
        styles['hebrew'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        styles['hebrew'].paragraph_format.keep_with_next = True
        styles['hebrew'].paragraph_format.space_after = 0

    def _get_clause_node(self, nid: NodeIdentifier) -> int:
        """Assign a clause node for a given node."""
        clause_rank = self.tf_api.Nodes.otypeRank['clause']
        rank = self.tf_api.Nodes.otypeRank[nid.otype]
        if rank > clause_rank:
            raise Exception(f'{nid} has a otype > clause!')
        return self.tf_api.L.u(nid.oslots[0], 'clause')[0]

    def _cluster_labels_by_clause(self, labels: List[LingLabel]):
        """Cluster labels by clause."""
        cl_clustered_labels = collections.defaultdict(list)
        for label in labels:
            cl_node = self._get_clause_node(label.nid)
            cl_clustered_labels[cl_node].append(label)
        return cl_clustered_labels

    def _add_reference_header(self, doc: Document, clause: int):
        """Add reference header to each entry."""
        en_book, ch, vs = self.tf_api.T.sectionFromNode(clause)
        ref = f'{en_book} {ch}:{vs}'
        book_node = self.tf_api.L.u(clause, 'book')[0]
        latin_book = self.tf_api.F.book.v(book_node)  # needed for SHEBANQ link
        shebanq_link = SHEBANQ_LINK.format(
            book=latin_book, chapter=str(ch), verse=str(vs)
        )
        heading = doc.add_paragraph(style=self.styles['ref'].name)
        add_hyperlink(heading, shebanq_link, ref)

    def _get_verse_nodes_from_clause(self, clause: int):
        """
        Get verse nodes from a clause node.

        Some clauses can span verses. So standard tf_api.L.u(clause, 'verse')
        will not work. We need to go from the slot level to get the one-to-many
        possible clause-verse relations.
        """
        verses = set()
        for word in self.tf_api.L.d(clause, 'word'):
            verses.add(self.tf_api.L.u(word, 'verse')[0])
        return sorted(verses)

    def _add_verse_text(self, doc: Document, clause: int):
        """Add verse hebrew text."""
        verse_nodes = self._get_verse_nodes_from_clause(clause)
        verse_text = doc.add_paragraph(
            self.tf_api.T.text(verse_nodes),
            style=self.styles['hebrew'].name,
        )
        verse_text.runs[0].font.color.rgb = GRAY

    def _add_clause_text(self, doc: Document, clause: int):
        """Add clause Hebrew text."""
        text = self.tf_api.T.text(clause)
        doc.add_paragraph(text, style=self.styles['hebrew'].name)

    def _add_slot_id_text(self, doc: Document, clause: int):
        """Add Hebrew text with slot identifiers indicated."""
        node_text = doc.add_paragraph(style=self.styles['hebrew'].name)
        for word in self.tf_api.L.d(clause, 'word'):
            node = node_text.add_run(str(word))
            node.font.superscript = True
            node.font.rtl = True
            node.font.color.rgb = RED
            heb_word = node_text.add_run(self.tf_api.T.text(word))
            heb_word.font.rtl = True
            heb_word.font.name = 'Times New Roman'
            heb_word.font.color.rgb = GRAY

    @staticmethod
    def _fix_autofit_bug(table):
        """
        Fix autofit bug for tables.

        Source:
        https://github.com/python-openxml/python-docx/issues/209#issuecomment-344417132
        """
        for col in table.columns:
            for cell in col.cells:
                cell._tc.tcPr.tcW.type = 'auto'

    def _get_label_node_text(self, label: LingLabel):
        """Get text for a label obj."""
        big_nodes = {
            'clause', 'sentence', 'verse', 'chapter', 'book'
        }
        if label.nid.otype in big_nodes:
            return f'[{label.nid.otype}]'
        else:
            return self.tf_api.T.text(label.nid.oslots)

    def _sort_labels(self, label: LingLabel) -> int:
        """Get sorting value for label."""
        return self.project.LABEL_ORDER[label.label]

    def _add_annotation_table(self, doc: Document, labels: List[LingLabel]):
        """Add annotation table to the document."""
        N_COLUMNS = 5
        table = doc.add_table(rows=0, cols=N_COLUMNS, style='Table Grid')
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.style.font.name = 'Helvetica Neue'
        table.style.paragraph_format.keep_with_next = True
        table.autofit = True
        for label in sorted(labels, key=self._sort_labels):
            row_cells = table.add_row().cells
            node_text = self._get_label_node_text(label)
            label_id = self.label_to_id[label]
            annotation_row = (
                label_id,
                label.label,
                label.target,
                node_text,
                label.value,
            )
            for (cell, text) in zip(row_cells, annotation_row):
                cell.text = text
            # adjust id cell to small size
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
            row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        self._fix_autofit_bug(table)

    def _build_document(self, document: Document()) -> None:
        """Build a document."""
        annotation_clusters = self._cluster_labels_by_clause(self.annotations)
        for clause in sorted(annotation_clusters):
            labels = annotation_clusters[clause]
            self._add_reference_header(document, clause)
            self._add_clause_text(document, clause)
            self._add_slot_id_text(document, clause)
            self._add_verse_text(document, clause)
            self._add_annotation_table(document, labels)
            document.add_paragraph('\n')

    @staticmethod
    def _label_from_row(row, metadata: Dict[str, Any]) -> LingLabel:
        """Extract a label from a row."""
        normalize_annotated_text = (
            lambda text: text.lower().strip()
        )
        try:
            id_cell, label_cell, target_cell, text_cell, value_cell = [
                normalize_annotated_text(cell.text) for cell in row.cells
            ]
        except ValueError:
            row_text = [cell.text for cell in row.cells]
            raise Exception(f'Misshapened row encountered: {row_text}')
        nid_data = metadata[id_cell]["nid"]
        nid = NodeIdentifier(nid_data[0], tuple(nid_data[1]))
        return LingLabel(
            label=label_cell,
            value=value_cell,
            nid=nid,
            target=target_cell,
        )
